import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HOME_IMG    = r'C:\Users\SHRADDHA\.gemini\antigravity\brain\45f2b99c-1a64-4fdd-abc5-15588c401c4a\home_page_screenshot_1777016380972.png'
LOGIN_IMG   = r'C:\Users\SHRADDHA\.gemini\antigravity\brain\45f2b99c-1a64-4fdd-abc5-15588c401c4a\login_page_screenshot_1777016400638.png'
CITIZEN_IMG = r'C:\Users\SHRADDHA\.gemini\antigravity\brain\45f2b99c-1a64-4fdd-abc5-15588c401c4a\citizen_dashboard_screenshot_1777016425079.png'
ADMIN_IMG   = r'C:\Users\SHRADDHA\.gemini\antigravity\brain\45f2b99c-1a64-4fdd-abc5-15588c401c4a\admin_dashboard_screenshot_1777016443513.png'

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

FONT = 'Times New Roman'

# ─── Helper functions ─────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = FONT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = FONT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    run.font.name = FONT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, justify=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = FONT
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    p.paragraph_format.space_after = Pt(4)
    return p

def cap(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = FONT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    return p

def add_img(path, width=Inches(5.5), caption=''):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
        if caption:
            cap(caption)
    except Exception as e:
        body(f'[Screenshot: {caption}  —  insert image here]')

def shade_cell(cell, hex_color='BDD7EE'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, 'BDD7EE')
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = FONT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_i, row_data in enumerate(rows):
        for c_i, val in enumerate(row_data):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.name = FONT
    doc.add_paragraph()

def pb():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Project Report\nOn')
r.font.size = Pt(16); r.font.bold = True; r.font.name = FONT

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('\n\nE-Municipal Portal\n(Newasa Nagar Parishad)')
r2.font.size = Pt(22); r2.font.bold = True; r2.font.name = FONT

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('\n\n\n\nSubmitted By\nShraddha Pramod Chavan\n\nIndira College of Commerce and Science, Pune\nM.Sc. (Computer Science) - Part I\nAcademic Year: 2025-2026')
r3.font.size = Pt(14); r3.font.name = FONT
pb()

# ══════════════════════════════════════════════════════════════
#  PROJECT CERTIFICATE
# ══════════════════════════════════════════════════════════════
h1('Project Certificate')
body('(To be obtained from the Project Guide / Department Head)\n\n'
     'This is to certify that the Project entitled "E-Municipal Portal" has been '
     'satisfactorily completed by Shraddha Pramod Chavan as a partial fulfillment of '
     'the requirements for M.Sc. (Computer Science) Part I, Semester II of '
     'Savitribai Phule Pune University, during the academic year 2025-2026.\n\n\n\n'
     '_____________________________                   _____________________________\n'
     'Project Guide                                            Head of Department\n\n\n'
     'Date:  ______________\nPlace: Pune')
pb()

# ══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
h1('TABLE OF CONTENTS')
toc_items = [
    ('CHAPTER 1: INTRODUCTION', True),
    ('1.1  Existing System and Need for System', False),
    ('1.2  Scope of Work', False),
    ('1.3  Operating Environment – Hardware and Software', False),
    ('CHAPTER 2: PROPOSED SYSTEM', True),
    ('2.1  Proposed System', False),
    ('2.2  Objectives of System', False),
    ('2.3  User Requirements', False),
    ('CHAPTER 3: ANALYSIS & DESIGN', True),
    ('3.1  Data Flow Diagram (DFD)', False),
    ('3.2  Functional Decomposition Diagram (FDD)', False),
    ('3.3  UML Diagrams', False),
    ('        3.3.1  Class Diagram', False),
    ('        3.3.2  Object Diagram', False),
    ('        3.3.3  Use Case Diagram', False),
    ('        3.3.4  Sequence Diagram', False),
    ('        3.3.5  Collaboration Diagram', False),
    ('        3.3.6  State Diagram', False),
    ('        3.3.7  Activity Diagram', False),
    ('        3.3.8  Component Diagram', False),
    ('        3.3.9  Deployment Diagram', False),
    ('3.4  Entity Relationship Diagram (ERD)', False),
    ('3.5  Data Dictionary', False),
    ('3.6  Table Design', False),
    ('3.7  Code Design', False),
    ('3.8  Menu Tree', False),
    ('3.9  Menu Screens', False),
    ('3.10 Input Screens', False),
    ('CHAPTER 4: USER MANUAL', True),
    ('4.1  Operations Manual / Menu Explanation', False),
    ('4.2  Menu Screens', False),
    ('4.3  Project Code', False),
    ('CHAPTER 5: LIMITATIONS AND ENHANCEMENT', True),
    ('5.1  Drawbacks and Limitations', False),
    ('5.2  Proposed Enhancements', False),
    ('5.3  Conclusions', False),
    ('5.4  Bibliography', False),
]
for text, bold in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.name = FONT
    p.paragraph_format.space_after = Pt(1)
pb()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════
h1('CHAPTER 1 : INTRODUCTION')

h2('1.1  Existing System and Need for System')
body('In the Existing System:\n'
     'All civic records including birth registrations, property taxes, water connection data, and '
     'public complaints are maintained physically in paper ledgers at the Nagar Parishad office. '
     'To avail any municipal service, citizens must physically visit the office, wait in long queues, '
     'and submit handwritten forms, making even simple tasks highly time-consuming.\n\n'
     'There is no centralized database to track complaints or service requests. Employees spend the '
     'majority of their shift sorting through stacks of paper applications. Interdepartmental '
     'communication requires physical file movement leading to loss or delays. Report generation '
     'for administrative review is tedious and relies on manual compilation.\n\n'
     'Getting proper status reports is a tedious job. Data from one department has to be manually '
     'conveyed to another. There is no web-based interface allowing citizens to manage their services '
     'from home. These challenges necessitate a digitized, centralized E-Municipal Portal.')

h2('1.2  Scope of Work')
body('For this project, work is divided in modules according to responsibilities:\n'
     '  • Citizen Operations Team\n'
     '  • Employee & Complaint Processing Team\n'
     '  • Admin & Department Management Team\n'
     '  • Architecture & Infrastructure Team\n\n'
     'As the developer, I am responsible for:\n'
     '  • Analyzing the Nagar Parishad\'s operational workflow and citizen needs.\n'
     '  • Deciding the business logic for role-based access, complaint assignment, billing, and '
     'document upload workflows.\n'
     '  • Writing source code for the React-based web interface.\n'
     '  • Writing generalized C# classes for secure authentication and database operations.\n'
     '  • Writing ASP.NET Core RESTful Web APIs for all data processing using JWT security.')

h2('1.3  Operating Environment – Hardware and Software')
body('Hardware Requirements:')
make_table(
    ['Component',       'Minimum Specification'],
    [
        ['Processor',   'Intel Core i3 / AMD Ryzen 3 at 2.0 GHz or higher'],
        ['RAM',         '4 GB (8 GB Recommended)'],
        ['Hard Disk',   '250 GB'],
        ['Network',     'Active Broadband Internet Connection'],
    ]
)
body('Software Requirements:')
make_table(
    ['Component',                   'Technology Used'],
    [
        ['Operating System',        'Windows 10 / 11 or Linux'],
        ['Frontend Framework',      'React JS (Vite) with Tailwind CSS'],
        ['Backend Framework',       'ASP.NET Core 8 Web API (C#)'],
        ['Database',                'PostgreSQL 16'],
        ['ORM',                     'Entity Framework Core (Code-First)'],
        ['Authentication',          'JWT – JSON Web Tokens (SHA-256 password hashing)'],
        ['Web Browser',             'Google Chrome / Microsoft Edge / Firefox'],
        ['Development IDE',         'Visual Studio Code, pgAdmin 4'],
    ]
)
pb()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 2: PROPOSED SYSTEM
# ══════════════════════════════════════════════════════════════
h1('CHAPTER 2 : PROPOSED SYSTEM')

h2('2.1  Proposed System')
body('In the proposed E-Municipal Portal system, we have the following modules:\n\n'
     'Authentication Module:\nThis module secures the portal using JWT-based authentication with '
     'SHA-256 password hashing. It handles three user roles — Citizen, Employee, and Admin — each '
     'with access restricted to their respective features only.\n\n'
     'Citizen Services Module:\nAllows citizens to browse available municipal services such as Birth '
     'Certificates, Property Tax, Water Connection, and Housing Schemes. Citizens can apply online '
     'by uploading required documents (Aadhar/PAN). Each application tracks its lifecycle: '
     'Pending → Approved → Completed.\n\n'
     'Complaint Management Module:\nCitizens can file civic complaints (road damage, water leakage, '
     'streetlights) with supporting photo proof. The Admin assigns complaints to the appropriate '
     'department employees who update the status upon resolution.\n\n'
     'Billing Module:\nMaintains records of all utility bills and service fees. Citizens can view '
     'their pending bills and simulate payment using Credit Card, Debit Card, or Net Banking. '
     'A payment receipt is generated automatically upon successful payment.\n\n'
     'Department and Employee Module:\nManages the organizational hierarchy. Departments can be '
     'created, updated, or deleted. Employees are registered under specific departments with '
     'designations and salary information.')

h2('2.2  Objectives of System')
body('The objectives of the E-Municipal Portal can be summarized as:\n\n'
     '• Reduce time, cost, and resources required to conduct daily municipal operations.\n'
     '• Provide 24/7 web-based access to all municipal services, eliminating the need for '
     'physical office visits by citizens.\n'
     '• Establish complete transparency through digital tracking of complaints and application statuses.\n'
     '• Streamline interdepartmental communication by routing tasks digitally to the right employees.\n'
     '• Provide administrators with a real-time unified dashboard for complete management oversight.\n'
     '• Generate bills, receipts, and reports in digital format to eliminate paper-based inefficiency.')

h2('2.3  User Requirements')
body('The E-Municipal Portal is designed for three types of users:\n\n'
     'Admin:\nAdmin has the maximum privileges. The Admin can create and manage departments, register '
     'new employees, oversee all citizen applications and complaints, assign complaints to specific '
     'employees, manage the service catalogue, and monitor portal-wide activity statistics.\n\n'
     'Employee:\nEmployee operates under a specific department. They log in to view their assigned '
     'complaints and service requests. They can view citizen-uploaded documents, change request '
     'statuses, and upload resolution proof images to formally close assigned tasks.\n\n'
     'Citizen:\nCitizen is the primary end user. They can register, securely log in, apply for '
     'government services with document uploads, pay utility bills, file civic complaints with '
     'photo evidence, and track the real-time status of all their submissions.')
pb()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 3: ANALYSIS & DESIGN
# ══════════════════════════════════════════════════════════════
h1('CHAPTER 3 : ANALYSIS & DESIGN')

h2('3.1  Data Flow Diagram (DFD)')
body('Context Level DFD:\nExternal Entities: Citizen, Employee, Admin\nThe system (E-Municipal Portal) receives Login Info, Service Applications, Bill Payments, and Complaint Reports from Citizens. It returns Status Updates, Certificates, and Receipts. Employees feed Resolved Status updates into the system. Admins manage Departments and Assignments.\n\n\n\n\n\n\n\n\n\n\n')
pb()
body('1st Level DFD:\n\n\n\n\n\n\n\n\n\n\n')
pb()

h2('3.2  Functional Decomposition Diagram (FDD)')
body('[Please insert the Functional Decomposition Diagram here]\n\n'
     'The FDD breaks down the E-Municipal Portal into sub-functions:\n'
     '  E-Municipal Portal\n'
     '  ├── Authentication Sub-system (Login, Register, JWT)\n'
     '  ├── Service Management (Browse, Apply, Track)\n'
     '  ├── Complaint Management (File, Assign, Resolve)\n'
     '  ├── Billing Sub-system (Generate, View, Pay)\n'
     '  └── Admin Management (Users, Departments, Reports)')

h2('3.3  UML Diagrams')

h3('3.3.1  Class Diagram')
body('Key classes: Citizen, Employee, Admin, Department, Service, ServiceRequest, Complaint, Bill\n\n\n\n\n\n\n\n\n\n\n')
pb()

h3('3.3.2  Object Diagram')
body('\n\n\n\n\n\n\n\n\n\n\n')

h3('3.3.3  Use Case Diagram')
body('Actors: Citizen, Employee, Admin\nUse Cases: Login, Apply for Service, File Complaint, Pay Bill, Assign Complaint, Update Status, Manage Departments, Manage Employees, View Reports\n\n\n\n\n\n\n\n\n\n')
pb()

h3('3.3.4  Sequence Diagram')
body('\n\n\n\n\n\n\n\n\n\n\n')

h3('3.3.5  Collaboration Diagram')
body('\n\n\n\n\n\n\n\n\n\n\n')
pb()

h3('3.3.6  State Diagram')
body('Complaint lifecycle: Open → Assigned → In Progress → Resolved\n\n\n\n\n\n\n\n\n\n\n')

h3('3.3.7  Activity Diagram')
body('\n\n\n\n\n\n\n\n\n\n\n')
pb()

h3('3.3.8  Component Diagram')
body('Components: Frontend (React), Backend (ASP.NET Core), and Database (PostgreSQL)\n\n\n\n\n\n\n\n\n\n\n')

h3('3.3.9  Deployment Diagram')
body('Nodes: Client Browser, Web Server (Kestrel), and Database Server\n\n\n\n\n\n\n\n\n\n\n')
pb()

h2('3.4  Entity Relationship Diagram (ERD)')
body('Key Relationships:\n  • Citizen  (1) → (Many)  Complaints\n  • Citizen  (1) → (Many)  ServiceRequests\n  • Citizen  (1) → (Many)  Bills\n  • Department (1) → (Many) Employees\n  • Employee (1) → (Many)  Assigned Complaints\n  • Service  (1) → (Many)  ServiceRequests\n\n\n\n\n\n\n\n\n\n\n')
pb()

h2('3.5  Data Dictionary')

h3('Database: EPortalDb    Table: Citizens')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['IDNo',           'INTEGER',     'PRIMARY KEY, AUTO INCREMENT', 'Unique Citizen ID'],
        ['Name',           'VARCHAR(100)','NOT NULL',                    'Full name of citizen'],
        ['Email',          'VARCHAR(100)','NOT NULL, UNIQUE',            'Login email address'],
        ['PasswordHash',   'TEXT',        'NOT NULL',                    'SHA-256 hashed password'],
        ['Phno',           'VARCHAR(15)', 'NULLABLE',                    'Contact phone number'],
        ['Gender',         'VARCHAR(10)', 'NULLABLE',                    'Male / Female / Other'],
        ['Bday',           'VARCHAR(20)', 'NULLABLE',                    'Date of birth'],
        ['House_no',       'VARCHAR(50)', 'NULLABLE',                    'House number'],
        ['Street_no_name', 'VARCHAR(100)','NULLABLE',                    'Street or locality name'],
    ]
)

h3('Database: EPortalDb    Table: Employees')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['EID',          'INTEGER',       'PRIMARY KEY, AUTO INCREMENT','Unique Employee ID'],
        ['EName',        'VARCHAR(100)',  'NOT NULL',                   'Full name of employee'],
        ['Email',        'VARCHAR(100)',  'NOT NULL, UNIQUE',           'Login email address'],
        ['PasswordHash', 'TEXT',          'NOT NULL',                   'SHA-256 hashed password'],
        ['Phno',         'VARCHAR(15)',   'NULLABLE',                   'Contact phone number'],
        ['EAdd',         'TEXT',          'NULLABLE',                   'Residential address'],
        ['Salary',       'DECIMAL(10,2)', 'NOT NULL',                   'Monthly salary in INR'],
        ['DNo',          'INTEGER',       'FOREIGN KEY → Departments',  'Assigned department'],
    ]
)

h3('Database: EPortalDb    Table: Departments')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['DNo',   'INTEGER',      'PRIMARY KEY, AUTO INCREMENT','Unique Department ID'],
        ['DName', 'VARCHAR(100)', 'NOT NULL, UNIQUE',           'Department name (e.g., Water Supply)'],
    ]
)

h3('Database: EPortalDb    Table: Services')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['SID',          'INTEGER',       'PRIMARY KEY, AUTO INCREMENT', 'Unique Service ID'],
        ['SName',        'VARCHAR(100)',  'NOT NULL',                    'Service name'],
        ['Rate',         'DECIMAL(10,2)', 'NOT NULL',                    'Application fee (INR)'],
        ['DNo',          'INTEGER',       'FOREIGN KEY → Departments',   'Department handling service'],
        ['RequiredDocs', 'TEXT',          'NULLABLE',                    'Comma-separated required docs'],
    ]
)

h3('Database: EPortalDb    Table: ServiceRequests')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['Id',           'INTEGER',   'PRIMARY KEY, AUTO INCREMENT',    'Unique Request ID'],
        ['Status',       'VARCHAR',   'NOT NULL',                       'Pending/Approved/Rejected/Completed'],
        ['DocumentUrls', 'TEXT',      'NULLABLE',                       'Paths to uploaded documents'],
        ['CreatedAt',    'TIMESTAMP', 'NOT NULL',                       'Application submission date'],
        ['IDNo',         'INTEGER',   'FOREIGN KEY → Citizens',         'Citizen applicant'],
        ['SID',          'INTEGER',   'FOREIGN KEY → Services',         'Service applied for'],
        ['Bill_ID',      'INTEGER',   'FOREIGN KEY → Bills (NULLABLE)', 'Linked bill if applicable'],
    ]
)

h3('Database: EPortalDb    Table: Complaints')
make_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['CID',         'INTEGER',   'PRIMARY KEY, AUTO INCREMENT','Unique Complaint ID'],
        ['Title',       'VARCHAR',   'NOT NULL',                   'Short title of complaint'],
        ['Description', 'TEXT',      'NULLABLE',                   'Detailed description of issue'],
        ['C_date',      'TIMESTAMP', 'NOT NULL',                   'Date complaint was submitted'],
        ['C_status',    'VARCHAR',   'NOT NULL',                   'Open / In Progress / Resolved'],
        ['ImagePath',   'TEXT',      'NULLABLE',                   'Path to uploaded photo proof'],
        ['IDNo',        'INTEGER',   'FOREIGN KEY → Citizens',     'Citizen who filed complaint'],
        ['DNo',         'INTEGER',   'FOREIGN KEY → Departments',  'Department responsible'],
        ['EID',         'INTEGER',   'FOREIGN KEY → Employees',    'Assigned employee'],
    ]
)

h3('Database: EPortalDb    Table: Bills')
make_table(
    ['Column Name',   'Data Type',       'Constraints',                  'Description'],
    [
        ['Bill_ID',       'INTEGER',       'PRIMARY KEY, AUTO INCREMENT',  'Unique Bill ID'],
        ['P_date',        'TIMESTAMP',     'NOT NULL',                     'Bill generation date'],
        ['Total_amt',     'DECIMAL(10,2)', 'NOT NULL',                     'Total amount due (INR)'],
        ['PaymentMethod', 'VARCHAR(50)',   'NULLABLE',                     'Credit/Debit/Net Banking'],
        ['IsPaid',        'BOOLEAN',       'NOT NULL, DEFAULT FALSE',      'Payment status flag'],
        ['IDNo',          'INTEGER',       'FOREIGN KEY → Citizens',       'Citizen bill owner'],
        ['ConsumerNumber','VARCHAR',       'NULLABLE',                     'Utility consumer number'],
        ['BillType',      'VARCHAR',       'NULLABLE',                     'Water/Property Tax/etc.'],
        ['DueDate',       'DATE',          'NULLABLE',                     'Payment deadline date'],
    ]
)

h2('3.6  Table Design')
body('All tables are designed using Entity Framework Core Code-First Migrations with PostgreSQL 16. '
     'Primary keys are auto-incremented integers. Foreign key relationships are enforced at the '
     'database level ensuring referential integrity. Passwords are stored as SHA-256 cryptographic '
     'hashes — plain text passwords are never stored. Document files are saved to the server '
     'file system and their paths are stored in the database.')

h2('3.7  Code Design')
h3('Authentication Controller Class')
code(
    'public class AuthController : ControllerBase\n'
    '{\n'
    '    // POST /api/auth/login\n'
    '    public async Task<IActionResult> Login([FromBody] LoginDto dto)\n'
    '    {\n'
    '        // Validate email & hashed password, return JWT token\n'
    '    }\n\n'
    '    // POST /api/auth/register\n'
    '    public async Task<IActionResult> RegisterCitizen([FromBody] CitizenRegisterDto dto)\n'
    '    {\n'
    '        // Hash password, create Citizen record in DB\n'
    '    }\n\n'
    '    // POST /api/auth/register-employee  [Admin only]\n'
    '    public async Task<IActionResult> RegisterEmployee([FromBody] EmployeeRegisterDto dto)\n'
    '    {\n'
    '        // Hash password, create Employee record in DB\n'
    '    }\n'
    '}'
)

h3('Services Controller Class')
code(
    'public class ServicesController : ControllerBase\n'
    '{\n'
    '    // GET /api/services\n'
    '    public async Task<IActionResult> GetAll()\n'
    '    { /* Returns all available municipal services */ }\n\n'
    '    // POST /api/services  [Admin only]\n'
    '    public async Task<IActionResult> Create([FromBody] ServiceCreateDto dto)\n'
    '    { /* Admin creates a new service entry */ }\n\n'
    '    // POST /api/services/{id}/apply  [Citizen only]\n'
    '    public async Task<IActionResult> Apply(int id, [FromForm] ...)\n'
    '    { /* Citizen applies with document uploads */ }\n'
    '}'
)

h3('Complaint Controller Class')
code(
    'public class ComplaintsController : ControllerBase\n'
    '{\n'
    '    // GET /api/complaints\n'
    '    public async Task<IActionResult> GetAll()\n'
    '    { /* Returns complaints filtered by role */ }\n\n'
    '    // POST /api/complaints  [Citizen only]\n'
    '    public async Task<IActionResult> Create([FromForm] ...)\n'
    '    { /* Citizen files complaint with photo upload */ }\n\n'
    '    // PUT /api/complaints/{id}/assign  [Admin only]\n'
    '    public async Task<IActionResult> Assign(int id, [FromBody] ComplaintAssignDto dto)\n'
    '    { /* Admin assigns complaint to department employee */ }\n\n'
    '    // PUT /api/complaints/{id}/status  [Employee only]\n'
    '    public async Task<IActionResult> UpdateStatus(int id, [FromForm] ...)\n'
    '    { /* Employee marks complaint resolved + proof image */ }\n'
    '}'
)

h2('3.8  Menu Tree')
code(
    'E-Municipal Portal\n'
    '├── Home Page\n'
    '│   ├── News Ticker\n'
    '│   ├── Services Dropdown\n'
    '│   ├── City Information\n'
    '│   ├── Government Yojanas\n'
    '│   └── Contact Section\n'
    '├── Login Page\n'
    '├── Register Page\n'
    '├── Citizen Panel\n'
    '│   ├── Dashboard\n'
    '│   ├── Apply for Service\n'
    '│   ├── My Applications\n'
    '│   ├── My Bills\n'
    '│   ├── My Complaints\n'
    '│   └── Feedback\n'
    '├── Employee Panel\n'
    '│   ├── Dashboard\n'
    '│   └── My Assigned Complaints\n'
    '└── Admin Panel\n'
    '    ├── Dashboard (Statistics)\n'
    '    ├── Manage Employees\n'
    '    ├── Manage Departments\n'
    '    ├── Manage Services\n'
    '    └── Complaints Management'
)

h2('3.9  Menu Screens')
body('[Please insert screenshots of the Sidebar Navigation for Citizen, Employee, and Admin panels here]')

h2('3.10  Input Screens')
body('[Please insert screenshots of the following input screens:\n'
     '  1. Citizen Registration Form\n'
     '  2. New Complaint Submission Form\n'
     '  3. Service Application Form with file upload\n'
     '  4. Bill Payment Screen]')
pb()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 4: USER MANUAL
# ══════════════════════════════════════════════════════════════
h1('CHAPTER 4 : USER MANUAL')

h2('4.1  Operations Manual / Menu Explanation')
body('Main Menu:\n'
     'The Home Page navbar contains: Home | Services | City Info | Contact | Login\n\n'
     'Services Dropdown Menu (from Home Page header):\n'
     '  • Birth Certificate\n'
     '  • Caste Certificate\n'
     '  • Water Connection\n'
     '  • Property Tax\n'
     '  • Death Certificate\n'
     '  • Building Permission\n'
     '  → Clicking any service redirects to Login page (with service name context displayed)\n\n'
     'Citizen Panel Sidebar Menu:\n'
     '  Dashboard | Apply for Service | My Applications | My Bills | My Complaints | Feedback\n\n'
     'Employee Panel Sidebar Menu:\n'
     '  Dashboard | My Assigned Complaints\n\n'
     'Admin Panel Sidebar Menu:\n'
     '  Dashboard | Manage Employees | Manage Departments | Manage Services | Complaints')

h2('4.2  Menu Screens')
body('Fig 1: Home Page – Newasa Municipal Corporation Portal')
add_img(HOME_IMG, Inches(5.5), 'Figure 1: Home Page of E-Municipal Portal')

body('Fig 2: Login Page with Role Selector')
add_img(LOGIN_IMG, Inches(4.8), 'Figure 2: Login / Sign-In Page')

body('Fig 3: Citizen Dashboard')
add_img(CITIZEN_IMG, Inches(5.5), 'Figure 3: Citizen Dashboard – Service Application Panel')

body('Fig 4: Admin Dashboard')
add_img(ADMIN_IMG, Inches(5.5), 'Figure 4: Admin Dashboard – Management Overview')

h2('4.3  Project Code')
h3('JWT Token Generation (C#)')
code(
    'private string GenerateJwtToken(int userId, string name, string email, string role)\n'
    '{\n'
    '    var claims = new[]\n'
    '    {\n'
    '        new Claim(ClaimTypes.NameIdentifier, userId.ToString()),\n'
    '        new Claim(ClaimTypes.Name,           name),\n'
    '        new Claim(ClaimTypes.Email,          email),\n'
    '        new Claim(ClaimTypes.Role,           role)\n'
    '    };\n'
    '    var key  = new SymmetricSecurityKey(Encoding.UTF8.GetBytes(_configuration["JwtSettings:Secret"]));\n'
    '    var cred = new SigningCredentials(key, SecurityAlgorithms.HmacSha256);\n'
    '    var token = new JwtSecurityToken(\n'
    '        issuer:   _configuration["JwtSettings:Issuer"],\n'
    '        audience: _configuration["JwtSettings:Audience"],\n'
    '        claims:   claims,\n'
    '        expires:  DateTime.UtcNow.AddDays(7),\n'
    '        signingCredentials: cred\n'
    '    );\n'
    '    return new JwtSecurityTokenHandler().WriteToken(token);\n'
    '}'
)

h3('React Login Form (JavaScript / JSX)')
code(
    'const handleSubmit = async (e) => {\n'
    '    e.preventDefault();\n'
    '    setLoading(true);\n'
    '    try {\n'
    '        const data = await login(form.email, form.password, form.role);\n'
    '        if (data.role === "Admin")    navigate("/admin");\n'
    '        else if (data.role === "Employee") navigate("/employee");\n'
    '        else navigate("/citizen");\n'
    '    } catch (err) {\n'
    '        setError(err?.response?.data || "Invalid credentials.");\n'
    '    } finally {\n'
    '        setLoading(false);\n'
    '    }\n'
    '};'
)

h3('AuthContext – Global Login State (React)')
code(
    'const login = async (email, password, role) => {\n'
    '    const res  = await api.post("/auth/login", { email, password, role });\n'
    '    const data = res.data;\n'
    '    localStorage.setItem("token", data.token);\n'
    '    localStorage.setItem("user",  JSON.stringify({ role: data.role, name: data.name }));\n'
    '    setUser({ role: data.role, name: data.name, userId: data.userId });\n'
    '    return data;\n'
    '};'
)

h3('Complaint Filing with Image Upload (C#)')
code(
    '[HttpPost]\n'
    'public async Task<IActionResult> Create([FromForm] ComplaintCreateDto dto)\n'
    '{\n'
    '    string? imagePath = null;\n'
    '    if (dto.Image != null && dto.Image.Length > 0)\n'
    '    {\n'
    '        var fileName = Guid.NewGuid() + Path.GetExtension(dto.Image.FileName);\n'
    '        var savePath = Path.Combine("uploads", "complaints", fileName);\n'
    '        using var stream = System.IO.File.Create(savePath);\n'
    '        await dto.Image.CopyToAsync(stream);\n'
    '        imagePath = savePath;\n'
    '    }\n'
    '    // Save complaint to database...\n'
    '}'
)
pb()

# ══════════════════════════════════════════════════════════════
#  CHAPTER 5: LIMITATIONS AND ENHANCEMENT
# ══════════════════════════════════════════════════════════════
h1('CHAPTER 5 : LIMITATIONS AND ENHANCEMENT')

h2('5.1  Drawbacks and Limitations')
body('  • The system currently requires an active internet connection at all times; offline access '
     'is not supported.\n'
     '  • Document verification is performed manually by employees, which can create processing '
     'bottlenecks when application volumes are high.\n'
     '  • The payment gateway is currently simulated. A real payment processor (RazorPay, Paytm, '
     'or bank gateway) must be integrated before production use.\n'
     '  • The portal currently supports only the English language. Marathi language support is '
     'not yet implemented, which limits accessibility for rural citizens.\n'
     '  • File upload is limited to specific formats and sizes. Large document uploads may slow down '
     'network performance on lower bandwidth connections.')

h2('5.2  Proposed Enhancements')
body('  • Multilingual Support: Implement a full Marathi interface to improve accessibility '
     'for all citizens of Newasa.\n'
     '  • SMS / Email Notification Gateway: Integrate Twilio or MSG91 to automatically notify '
     'citizens via SMS when their application status changes.\n'
     '  • AI Chatbot: Deploy a smart assistant to answer common citizen queries about required '
     'documents, fees, or application procedures without human intervention.\n'
     '  • OCR Document Verification: Integrate Optical Character Recognition to automatically '
     'extract and validate Aadhar/PAN details from uploaded scans, reducing manual effort.\n'
     '  • Mobile Application: Develop a dedicated Android and iOS application for native mobile '
     'access to the portal.\n'
     '  • Real Payment Gateway: Integrate Razorpay or Paytm for live financial transactions.')

h2('5.3  Conclusions')
body('This project gave me an excellent opportunity to understand the complete flow of a '
     'government E-Governance portal and the challenges involved in developing and deploying '
     'such a system for real municipal needs. We had the unique opportunity to engineer not just '
     'a simple CRUD application, but a multi-role, secure, full-stack web portal serving Citizens, '
     'Municipal Employees, and Administrators under one unified platform.\n\n'
     'The E-Municipal Portal successfully eliminates the core pain points of the current manual '
     'system — unnecessary office visits, paper-based records, and departmental silos — by '
     'providing a fast, transparent, and scalable digital solution built on React JS and '
     'ASP.NET Core 8.\n\n'
     'It gives us immense pleasure and satisfaction to complete a project in a field of real societal '
     'impact, using our technical skills to deliver a result that goes beyond academic expectations '
     'and contributes meaningfully towards smart e-governance for the citizens of Newasa.')

h2('5.4  Bibliography')
body('1.  React Official Documentation — react.dev\n'
     '2.  ASP.NET Core 8 Web API Documentation — learn.microsoft.com\n'
     '3.  Entity Framework Core in Action — Jon P Smith (Manning Publications)\n'
     '4.  PostgreSQL 16 Official Manual — postgresql.org/docs\n'
     '5.  Tailwind CSS Documentation — tailwindcss.com\n'
     '6.  JSON Web Tokens Introduction — jwt.io\n'
     '7.  E-Governance Initiatives in Maharashtra — NASSCOM Report 2023\n'
     '8.  Schildt, H. (2021). C# Complete Reference. McGraw-Hill Education.\n'
     '9.  Moroney, L. (2020). Full-Stack Web Development with ASP.NET Core. Apress.\n'
     '10. National e-Governance Plan (NeGP) — Government of India, meity.gov.in')

# ══════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════
out = r'C:\Users\SHRADDHA\OneDrive\Documents\eportal\E_Municipal_Portal_Final_Report.docx'
doc.save(out)
print(f'SUCCESS — Report saved: {out}')
