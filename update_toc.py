import docx
from docx.shared import Pt

try:
    doc = docx.Document(r'C:\Users\SHRADDHA\OneDrive\Documents\eportal\E_Municipal_Portal_Final_Report.docx')

    toc_items = [
        ('CHAPTER 1: INTRODUCTION', True),
        ('1.1  Existing System and Need for System', False),
        ('1.2  Scope of Work', False),
        ('1.3  Operating Environment – Hardware and Software', False),
        ('CHAPTER 2: PROPOSED SYSTEM', True),
        ('2.1  Proposed System', False),
        ('2.2  Objectives of System', False),
        ('2.3  User Requirements', False),
        ('CHAPTER 3: ANALYSIS & DESIGN', True),
        ('3.1  UML Diagrams', False),
        ('        3.1.1  Class Diagram', False),
        ('        3.1.2  Object Diagram', False),
        ('        3.1.3  Use Case Diagram', False),
        ('        3.1.4  Sequence Diagram', False),
        ('        3.1.5  Collaboration Diagram', False),
        ('        3.1.6  State Diagram', False),
        ('        3.1.7  Activity Diagram', False),
        ('        3.1.8  Component Diagram', False),
        ('        3.1.9  Deployment Diagram', False),
        ('3.2  Entity Relationship Diagram (ERD)', False),
        ('3.3  Data Dictionary', False),
        ('3.4  Table Design', False),
        ('3.5  Code Design', False),
        ('3.6  Menu Tree', False),
        ('CHAPTER 4: USER MANUAL', True),
        ('4.1  Operations Manual / Menu Explanation', False),
        ('4.2  Input Screens', False),
        ('4.3  Project Code', False),
        ('CHAPTER 5: LIMITATIONS AND ENHANCEMENT', True),
        ('5.1  Drawbacks and Limitations', False),
        ('5.2  Proposed Enhancements', False),
        ('5.3  Conclusions', False),
        ('5.4  Bibliography', False),
    ]

    toc_start = -1
    toc_end = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == 'TABLE OF CONTENTS':
            toc_start = i
        elif text == 'CHAPTER 1 : INTRODUCTION' and toc_start != -1:
            toc_end = i
            break

    if toc_start != -1 and toc_end != -1:
        print(f"Found TOC delimiters at index {toc_start} and {toc_end}")
        # Delete old
        for i in range(toc_end - 1, toc_start, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        ref_p = doc.paragraphs[toc_start + 1] 
        for text, is_bold in toc_items:
            new_p = ref_p.insert_paragraph_before(text)
            if new_p.runs:
                run = new_p.runs[0]
                run.font.size = Pt(12)
                run.font.bold = is_bold
                run.font.name = 'Times New Roman'
            else:
                run = new_p.add_run()
                run.font.size = Pt(12)
                run.font.bold = is_bold
                run.font.name = 'Times New Roman'
            new_p.paragraph_format.space_after = Pt(1)

        # Fix headings inside body
        for p in doc.paragraphs:
            if 'Input Screens' in p.text and p.text.strip().startswith('3.10'):
                p.text = '4.2  Input Screens'
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                
        doc.save(r'C:\Users\SHRADDHA\OneDrive\Documents\eportal\E_Municipal_Portal_Final_Report.docx')
        print('TOC successfully updated!')
    else:
        print(f'Could not find TOC. toc_start={toc_start}, toc_end={toc_end}')

except Exception as e:
    print("Error:", e)
